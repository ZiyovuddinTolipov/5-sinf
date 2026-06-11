"""
Generate .docx lesson files for all subjects and save to admin/uploads/lessons/.
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("pip install python-docx")
    raise

OUTPUT_DIR = Path(__file__).parent / "admin" / "uploads" / "lessons"

LESSONS = {
    "Natural sonlar va amallar": {
        "subject": "Matematika",
        "sections": [
            ("Natural sonlar", "Natural sonlar — bu 1, 2, 3, 4, 5, ... kabi sanash uchun ishlatiladigan musbat butun sonlar. Nol (0) natural son hisoblanmaydi.\n\nNatural sonlar o'q ustida chap tomondan o'ngga qarab joylashadi. Har bir natural son oldingi sondan 1 ga katta."),
            ("Arifmetik amallar", "Natural sonlar ustida to'rtta asosiy amal bajariladi:\n\n• Qo'shish: 25 + 37 = 62\n• Ayirish: 100 - 43 = 57\n• Ko'paytirish: 12 × 8 = 96\n• Bo'lish: 144 ÷ 12 = 12\n\nAmallarni bajarish tartibi: avval qavs ichidagi amallar, keyin ko'paytirish va bo'lish, so'ng qo'shish va ayirish."),
            ("Razryad va sinf", "Ko'p xonali sonlarda har bir raqamning o'rni muhim:\n\n• Birlar razryadi: 1, 2, ..., 9\n• O'nlar razryadi: 10, 20, ..., 90\n• Yuzlar razryadi: 100, 200, ..., 900\n\nMisol: 4 738 = 4 ming + 7 yuz + 3 o'n + 8 birlik"),
        ],
        "terms": [("Natural son", "1 dan boshlanadigan butun musbat son"), ("Razryad", "Raqamning sondagi o'rni")],
        "questions": ["1 dan 100 gacha natural sonlar nechta?", "Qo'shishning almashtiruvchanlik xossasi nima?"],
    },
    "Kasrlar va o'nli kasrlar": {
        "subject": "Matematika",
        "sections": [
            ("Oddiy kasrlar", "Oddiy kasr — butun narsaning qismini ko'rsatadi. Kasrning ikki qismi bor:\n\n• Surat (ustdagi son) — nechtaданini ko'rsatadi\n• Mahraj (ostdagi son) — nechta teng qismga bo'linganini ko'rsatadi\n\nMisol: 3/4 — to'rtga bo'lingan butunning uchta qismi."),
            ("O'nli kasrlar", "O'nli kasr — o'nlik, yuzlik, minglik ulushlarni ifodalaydi.\n\nMisol: 0,75 = 75/100 = 3/4\n\nO'nli kasrlarni qo'shish: vergulni to'g'rilab qo'yib, oddiy sonlar kabi qo'shiladi.\n\n3,14 + 2,86 = 6,00"),
            ("Aralash son", "Aralash son = butun qism + kasr qism\n\nMisol: 2 3/4 = 2 + 3/4 = 11/4\n\nAralash sonni noto'g'ri kasrga aylantirish: butunni mahrajga ko'paytir, suratga qo'sh."),
        ],
        "terms": [("Kasr", "Butun narsaning qismi"), ("Surat", "Kasrning ustidagi soni"), ("Mahraj", "Kasrning ostidagi soni")],
        "questions": ["1/2 + 1/4 = ?", "0,5 ni oddiy kasrga aylantiring."],
    },
    "So'z turkumlari": {
        "subject": "Ona tili",
        "sections": [
            ("So'z turkumlari haqida", "O'zbek tilida so'zlar ma'no va grammatik xususiyatlariga ko'ra so'z turkumlariga bo'linadi:\n\n• Mustaqil so'z turkumlari: ot, sifat, son, olmosh, fe'l, ravish\n• Yordamchi so'z turkumlari: ko'makchi, bog'lovchi, yuklama\n• Undalma, undov, taqlid so'zlar"),
            ("Ot", "Ot — kishilarning ismi, narsalar, hodisalarning nomini bildiradi. Kimning? Nimaning? savollariga javob beradi.\n\nMisollar: kitob, uy, Alisher, maktab, baxt, sevgi\n\nOtlar songa (birlik/ko'plik), kelishikka o'zgaradi."),
            ("Sifat va Fe'l", "Sifat — predmetning belgisini bildiradi. Qanday? Qanaqa? savollariga javob beradi.\nMisol: katta, yashil, baland, chiroyli\n\nFe'l — harakat va holatni bildiradi. Nima qildi? Nima qilmoqda? savollariga javob beradi.\nMisol: o'qidi, yurmoqda, ishlamoqchi"),
        ],
        "terms": [("Ot", "Narsalarning nomini bildiruvchi so'z"), ("Sifat", "Predmet belgisini bildiruvchi so'z"), ("Fe'l", "Harakat va holatni bildiruvchi so'z")],
        "questions": ["\"Maktab\" so'zi qaysi turkum?", "Sifatga savol bering."],
    },
    "Gap va uning bo'laklari": {
        "subject": "Ona tili",
        "sections": [
            ("Gap va uning turlari", "Gap — tugallangan fikrni ifodalovchi so'z yoki so'zlar birikmasidir.\n\nGap turlari:\n• Darak gap: Bola maktabga ketdi.\n• So'roq gap: U qaerga ketdi?\n• Buyruq gap: Kitobni o'qi!\n• Undov gap: Qanday ajoyib!"),
            ("Gapning bosh bo'laklari", "Ega — gapda kim yoki nima haqida gapirilaётganini bildiradi. Kim? Nima? savollariga javob beradi.\n\nKesim — eganing ishi, harakati, holatini bildiradi. Nima qildi? Nima? savollariga javob beradi.\n\nMisol: Alisher kitob o'qidi.\n• Ega: Alisher\n• Kesim: o'qidi"),
            ("Ikkinchi darajali bo'laklar", "To'ldiruvchi — kimni? nimani? kimga? nimaga? savollariga javob beradi.\nAniqlovchi — qaysi? qanday? nechi? savollariga javob beradi.\nHol — qayerda? qachon? qanday? savollariga javob beradi."),
        ],
        "terms": [("Gap", "Tugallangan fikrni ifodalovchi birlik"), ("Ega", "Gapning asosiy bo'lagi — kim/nima"), ("Kesim", "Gapning asosiy bo'lagi — nima qildi")],
        "questions": ["Darak gapga misol keltiring.", "Gapning bosh bo'laklari qaysilar?"],
    },
    "Xalq og'zaki ijodi": {
        "subject": "Adabiyot",
        "sections": [
            ("Xalq og'zaki ijodi", "Xalq og'zaki ijodi — yozib qoldirilmagan, og'izdan-og'izga, avloddan-avlodga o'tib kelgan badiiy ijod namunalaridir.\n\nXalq og'zaki ijodiga quyidagilar kiradi:\n• Ertaklar\n• Dostonlar\n• Maqollar va topishmoqlar\n• Qo'shiqlar (lapar, yalla)\n• Rivoyat va afsonalar"),
            ("Ertaklar", "Ertak — xayoliy voqealarni tasvirlovchi badiiy asar. Ertaklar odatda baxtli yakunga ega bo'ladi.\n\nErtak turlari:\n• Hayvonlar haqidagi ertaklar\n• Sehrli ertaklar\n• Maishiy ertaklar\n\nO'zbek xalq ertaklari: \"Zumrad va Qimmat\", \"Oftob opa\", \"Yaxshilik qilsang, yaxshilik topasan\""),
            ("Maqol va topishmoqlar", "Maqol — qisqa, hikmatli, ko'p asrlik tajribani ifodalovchi ibora.\nMisol: \"Mehmonga ikki kun — uchinchi kun o'ziga ayon\"\n\nTopishmoq — ko'chma ma'noda biror narsani tasvirlab, uni topishni talab qiladigan asar.\nMisol: \"Oyoqsiz yuradi, tilsiz gapiradi\" — dарё"),
        ],
        "terms": [("Ertak", "Xayoliy voqealarni tasvirlovchi xalq ijodi"), ("Maqol", "Qisqa hikmatli xalq iborasi"), ("Topishmoq", "Ko'chma tasvirda biror narsani topish")],
        "questions": ["Ertakning asosiy xususiyati nima?", "Maqolga misol keltiring."],
    },
    "O'zbek klasssik adabiyoti": {
        "subject": "Adabiyot",
        "sections": [
            ("Alisher Navoiy", "Alisher Navoiy (1441–1501) — o'zbek adabiyotining eng buyuk vakili. U o'zbek adabiy tilini yaratdi va rivojlantirdi.\n\nAsosiy asarlari:\n• \"Xamsa\" — besh dostondan iborat\n• \"Xazoyin ul-maoniy\" — she'rlar to'plami\n• \"Muhokamat ul-lug'atayn\" — o'zbek va fors tillari taqqoslovi"),
            ("Zahiriddin Muhammad Bobur", "Bobur (1483–1530) — shoir, davlat arbobi va tarixchi.\n\nAsosiy asari: \"Boburnoma\" — o'z hayoti va davri haqidagi xotiralar kitobi. Bu asar jahon adabiyotining durdonalaridan biri hisoblanadi.\n\nBobur Hindistonda Boburiylar sulolasiga asos solgan."),
            ("O'zbek adabiyotining rivojlanishi", "XIX–XX asr o'zbek adabiyoti vakillari:\n• Muqimiy (1850–1903) — hajviy she'rlar muallifi\n• Furqat (1858–1909) — lirik shoir\n• Hamza Hakimzoda (1889–1929) — dramaturg va shoir\n• Abdulla Qodiriy (1894–1938) — romannavis, \"O'tkan kunlar\" muallifi"),
        ],
        "terms": [("Navoiy", "O'zbek adabiyotining buyuk vakili"), ("Boburnoma", "Boburning hayoti haqidagi xotiralar kitobi"), ("Xamsa", "Navoiyning besh dostondan iborat asari")],
        "questions": ["Navoiyning besh dostondan iborat asari nima deyiladi?", "Boburnoma qanday asar?"],
    },
    "Alphabet & Numbers": {
        "subject": "Ingliz tili",
        "sections": [
            ("English Alphabet", "English alphabet has 26 letters:\nA B C D E F G H I J K L M N O P Q R S T U V W X Y Z\n\nVowels (unli): A E I O U\nConsonants (undosh): all other 21 letters\n\nCapital letters — zaglav harf: A, B, C...\nSmall letters — kichik harf: a, b, c..."),
            ("Numbers 1–100", "1=one, 2=two, 3=three, 4=four, 5=five\n6=six, 7=seven, 8=eight, 9=nine, 10=ten\n11=eleven, 12=twelve, 13=thirteen, 14=fourteen, 15=fifteen\n20=twenty, 30=thirty, 40=forty, 50=fifty\n100=one hundred\n\n21=twenty-one, 35=thirty-five, 99=ninety-nine"),
            ("Basic Vocabulary", "Colors: red, blue, green, yellow, white, black, orange, purple\nAnimals: cat, dog, bird, fish, horse, cow, lion, tiger\nFamily: mother, father, sister, brother, grandmother, grandfather\nSchool: teacher, student, book, pen, desk, board, classroom"),
        ],
        "terms": [("Alphabet", "26 ta harfdan iborat ingliz alifbosi"), ("Vowel", "Unli tovush: A E I O U"), ("Number", "Son — one, two, three...")],
        "questions": ["How many letters are in English alphabet?", "What are vowels?"],
    },
    "Present Simple & Family": {
        "subject": "Ingliz tili",
        "sections": [
            ("Present Simple Tense", "Present Simple — hozirgi zamon fe'li. Doimiy holat yoki odatiy harakatlarni ifodalaydi.\n\nFormula:\nI/You/We/They + fe'l (asl holat)\nHe/She/It + fe'l + s/es\n\nMisol:\nI read books. — Men kitob o'qiyman.\nShe reads books. — U kitob o'qiydi.\nWe play football. — Biz futbol o'ynaymiz."),
            ("Negative and Question", "Negative (inkor): do not / does not + fe'l\nI do not like coffee. — Men qahva yoqtirmayman.\nHe does not play tennis.\n\nQuestion (so'roq): Do/Does + subject + fe'l?\nDo you speak English? — Siz inglizcha gaplashasizmi?\nDoes she work here?"),
            ("Family Members", "Mother — ona\nFather — ota\nSister — singil / opa\nBrother — uka / aka\nGrandmother — buvi / momо\nGrandfather — bobo\nAunt — xola / amaki\nUncle — tog'a / amaki\nCousin — amakivachcha\n\nMy family = mening oilam\nI have a sister and a brother."),
        ],
        "terms": [("Present Simple", "Hozirgi oddiy zamon"), ("Negative", "Inkor shakl"), ("Question", "So'roq shakl")],
        "questions": ["She ___ (play) tennis every day. To'g'ri shaklni qo'ying.", "\"Do not\" ning qisqartmasi nima?"],
    },
    "Qadimgi Sharq sivilizatsiyalari": {
        "subject": "Tarix",
        "sections": [
            ("Sivilizatsiya tushunchasi", "Sivilizatsiya — yuksak madaniyatga, yozuvga, shaharlarga, davlat tizimiga ega bo'lgan jamiyat. Dunyodagi dastlabki sivilizatsiyalar quyidagi hududlarda paydo bo'lgan:\n\n• Mesopotamiya (Dajla va Furot oralig'i)\n• Misr (Nil bo'yi)\n• Hindiston (Hind va Gang bo'yi)\n• Xitoy (Xuanxe va Yangtsze bo'yi)"),
            ("Qadimgi Misr sivilizatsiyasi", "Misr sivilizatsiyasi e.av. 3000-yildan boshlab rivojlandi. Asosiy yutuqlari:\n• Ehromlar (Xeops ehromining balandligi 147 metr)\n• Iyeroglif yozuvi\n• Papirus qog'ozi\n• Mumiyolash san'ati\n• 365 kunlik yil taqvimi"),
            ("Mesopotamiya sivilizatsiyasi", "Shumer va Bobil sivilizatsiyalari Mesopotamiyada rivojlandi:\n• Mixxat yozuvi — dunyodagi birinchi yozuv\n• G'ildirak ixtirosi (e.av. 3500)\n• Zikkurat — ibodatxona minorasi\n• Hammurapi qonunlari — 282 modda\n• Ashshurbanipal kutubxonasi — 30 000+ loy lavha"),
        ],
        "terms": [("Sivilizatsiya", "Yuksak madaniyatga ega bo'lgan jamiyat"), ("Iyeroglif", "Misr yozuvi"), ("Mixxat", "Shumer loy lavhaga yoziladigan yozuvi")],
        "questions": ["Birinchi yozuv qayerda paydo bo'lgan?", "Xeops ehromining balandligi necha metr?"],
    },
    "O'rta Osiyo tarixidan sahifalar": {
        "subject": "Tarix",
        "sections": [
            ("Zardushtiylik va Avesto", "Zardushtiylik — O'rta Osiyoda tarqalgan qadimgi din. Asosi: yaxshilik va yomonlik o'rtasidagi kurash.\nMuqaddas kitob — Avesto. Zardushtiylikning uch asosi:\n• Humata — Yaxshi fikr\n• Huxta — Yaxshi so'z\n• Huvarshta — Yaxshi ish"),
            ("Ipak yo'li", "Ipak yo'li — e.av. II asrdan eramizning XV–XVI asrlarigacha faoliyat ko'rsatgan savdo yo'li. Xitoy bilan G'arbiy Yevropa va Yaqin Sharqni bog'lagan.\n\nO'rta Osiyo shaharlari Ipak yo'lining markazida: Samarqand, Buxoro, Marv, Urganch.\n\nSavdo mahsulotlari: Xitoydan ipak, ziravorlar; G'arbdan shisha, qimmatbaho metallar."),
            ("Amir Temur davri", "Amir Temur (1336–1405) — Sohibqiron. U Movarounnahrda kuchli davlat barpo etdi. Paytaxti Samarqand.\n\nTemur davridagi inshootlar: Bibixonim masjidi, Go'ri Amir maqbarasi, Shohizinda ansambli.\n\nTimuriylar davri — madaniyat va ilm-fan gullagan davr. Ulug'bek (1394–1449) — astronom va matematik."),
        ],
        "terms": [("Avesto", "Zardushtiylikning muqaddas kitobi"), ("Ipak yo'li", "Xitoyni G'arb bilan bog'lagan savdo yo'li"), ("Sohibqiron", "Amir Temurning laqabi — 'Sayyoralarning egasi'")],
        "questions": ["Zardushtiylikning uch asosi nima?", "Amir Temurning paytaxti qaysi shahar edi?"],
    },
    "Yer shari va materiklar": {
        "subject": "Geografiya",
        "sections": [
            ("Yer sharining tuzilishi", "Yer — Quyosh sistemasidagi beshinchi eng katta sayyora. Yerni o'rganuvchi fan — geografiya.\n\nYer ichki tuzilishi:\n• Yadro — ichki (qattiq) va tashqi (suyuq)\n• Mantiya — qalin qatvar\n• Yer qobig'i — eng ustki qatlam\n\nYer yuzasining 71% suv, 29% quruqlik."),
            ("Materiklar", "Yer yuzida 6 ta materik (qit'a) bor:\n1. Osiyo — eng katta (44,6 mln km²)\n2. Afrika — ikkinchi katta\n3. Shimoliy Amerika\n4. Janubiy Amerika\n5. Antarktida — eng sovuq\n6. Avstraliya — eng kichik\n\nAvropa materik emas — Osiyo bilan birgalikda Evrosiyoni tashkil etadi."),
            ("Okeanlar", "Yer yuzida 4 ta okean:\n1. Tinch okeani — eng katta va chuqur\n2. Atlantika okeani — ikkinchi katta\n3. Hind okeani\n4. Arktika okeani — eng kichik\n\nTinch okeani yuzasi barcha quruqliklarning yuzidan kattaroq!"),
        ],
        "terms": [("Materik", "Okean bilan o'ralgan katta quruqlik"), ("Okean", "Eng katta suv havzasi"), ("Geografiya", "Yer shari haqidagi fan")],
        "questions": ["Yer yuzasining necha foizi suv?", "Eng katta materik qaysi?"],
    },
    "O'zbekiston tabiati va iqlimi": {
        "subject": "Geografiya",
        "sections": [
            ("O'zbekiston geografiyasi", "O'zbekiston — Markaziy Osiyodagi davlat. Maydoni — 448 900 km². Aholi — 37 million (2025 yil).\n\nChegaradosh davlatlar: Qozog'iston, Qirg'iziston, Tojikiston, Afg'oniston, Turkmaniston.\n\nPoytaxti: Toshkent — Markaziy Osiyoning eng yirik shahri."),
            ("Daryolar va ko'llar", "Asosiy daryolar:\n• Amudaryo — uzunligi 2 540 km, Orol dengiziga quyiladi\n• Sirdaryo — uzunligi 3 019 km\n• Zarafshon — Samarqand va Buxoro viloyatlaridan o'tadi\n• Farg'ona vodiysi daryolari\n\nKo'llar: Orol dengizi (qurib bormoqda), Aydarkul"),
            ("Iqlim va tabiat", "O'zbekiston iqlimi — keskin kontinental. Yoz issiq va quruq (40°C gacha), qish sovuq.\n\nTabiat mintaqalari:\n• Cho'l va yarim cho'l — Qizilqum cho'li\n• Tog' hududlari — Tyan-Shan, Pomir-Oloy\n• Vodiylar — Farg'ona, Zarafshon, Surxon\n\nTabiiy boyliklar: tabiiy gaz, neft, oltin, uran, paxta"),
        ],
        "terms": [("Amudaryo", "O'zbekistonning eng katta daryosi"), ("Qizilqum", "O'zbekistondagi katta cho'l"), ("Iqlim", "Hududning ko'p yillik ob-havo xususiyati")],
        "questions": ["O'zbekiston qancha davlat bilan chegaradosh?", "Qizilqum nima?"],
    },
    "O'simliklar dunyosi": {
        "subject": "Biologiya",
        "sections": [
            ("O'simliklar haqida", "O'simliklar — tirik organizmlarning muhim guruhidir. Ular quyosh nuri yordamida fotosintez orqali oziq yaratadi.\n\nO'simliklarning asosiy qismlari:\n• Ildiz — suvni va mineral moddalarni yutadi\n• Poya — suv va oziqni tashiydi\n• Barg — fotosintez amalga oshadi\n• Gul — ko'payish organi\n• Meva — urug'ni himoya qiladi"),
            ("Fotosintez", "Fotosintez — o'simliklar quyosh nuri, suv va CO₂ dan organik modda va kislorod ishlab chiqarish jarayoni.\n\nFormula: 6CO₂ + 6H₂O + nur → C₆H₁₂O₆ + 6O₂\n\nFotosintez faqat yashil o'simliklarda xlorofill tufayli sodir bo'ladi. Barglardagi klorofill — yashil rangli moddadir."),
            ("O'simliklar turlari", "O'simliklar guruhlari:\n• Guldorlar — eng ko'p tarqalgan (250 000+ tur)\n• Ignabargli daraxtlar — qarag'ay, archa, sadr\n• Paporotniklar — spera orqali ko'payadi\n• Moxlar — eng sodda o'simliklar\n• Suvo'tlar — suvda yashovchi\n\nO'zbekistonda 4 500 dan ortiq o'simlik turi mavjud."),
        ],
        "terms": [("Fotosintez", "Quyosh nuri yordamida oziq yaratish jarayoni"), ("Xlorofill", "Barglardagi yashil bo'yoq modda"), ("Ildiz", "Suvni yutuvchi o'simlik organi")],
        "questions": ["Fotosintez qaysi organda boradi?", "Xlorofill nima?"],
    },
    "Inson tanasi tuzilishi": {
        "subject": "Biologiya",
        "sections": [
            ("Inson tanasining tizimlaри", "Inson tanasi bir nechta tizimdan iborat:\n• Suyak-mushak tizimi — harakatni ta'minlaydi\n• Qon-tomir tizimi — qon va oziqlarni tashiydi\n• Nafas olish tizimi — kislorod va CO₂ almashinuvi\n• Hazm qilish tizimi — oziqni hazm qiladi\n• Asab tizimi — boshqaruv markazi"),
            ("Yurak va qon tizimi", "Yurak — qonni butun tanaga haydovchi mushak organi. Bir daqiqada 60–80 marta uradi.\n\nQon — qizil va oq qon tanachalari, trombotsitlar va qon plazmasidan iborat.\n\nQizil qon tanachalari (eritrotsitlar) — kislorod tashiydi. Ulardagi gemoglobin qonga qizil rang beradi.\n\nOq qon tanachalari (leykotsitlar) — kasalliklarga qarshi kurashadi."),
            ("Suyaklar va mushaklar", "Inson tanasida 206 ta suyak bor. Eng uzun suyak — son suyagi.\n\nSuyaklar vazifalari:\n• Tanani ushlab turadi\n• Ichki organlarni himoya qiladi\n• Qon ishlab chiqaradi (ko'mik)\n\nMushaklар suyaklarga birikib, harakatni ta'minlaydi. Inson tanasida 600 dan ortiq mushak bor."),
        ],
        "terms": [("Gemoglobin", "Qizil qon tanachasidagi kislorod tashuvchi modda"), ("Eritrotsit", "Qizil qon tanachasi"), ("Suyak", "Tananing tayanchini tashkil qiluvchi qattiq to'qima")],
        "questions": ["Inson tanasida nechta suyak bor?", "Yurak bir daqiqada necha marta uradi?"],
    },
    "Kompyuter qismlari va funksiyalari": {
        "subject": "Informatika",
        "sections": [
            ("Kompyuter tuzilishi", "Kompyuter ikki qismdan iborat:\n\n1. Apparatli qism (Hardware) — jismoniy qurilmalar:\n• Protsessor (CPU) — \"miya\", hisoblashlar bajaradi\n• Xotira (RAM) — operativ xotira\n• Qattiq disk (HDD/SSD) — doimiy xotira\n• Monitor — ko'rsatish qurilmasi\n• Klaviatura va sichqoncha — kiritish qurilmalari"),
            ("Dasturiy ta'minot", "2. Dasturiy ta'minot (Software) — dasturlar:\n\n• Operatsion tizim (OS): Windows, macOS, Linux\n• Ilovalar: Word, Excel, Chrome, Telegram\n• Antivirus: virus va zararli dasturlardan himoya\n\nOS — kompyuterning asosiy dasturi. U barcha qurilmalar va ilovalarni boshqaradi."),
            ("Ma'lumot o'lchov birliklari", "Kompyuterdagi ma'lumot bitlarda o'lchanadi:\n\n• 1 bit — 0 yoki 1\n• 8 bit = 1 bayt\n• 1 024 bayt = 1 kilobayt (KB)\n• 1 024 KB = 1 megabayt (MB)\n• 1 024 MB = 1 gigabayt (GB)\n\nOdatiy fayllar: rasm ~2 MB, musiqa ~5 MB, kino ~1.5 GB"),
        ],
        "terms": [("CPU", "Protsessor — kompyuterning hisoblash qismi"), ("RAM", "Operativ xotira"), ("OS", "Operatsion tizim")],
        "questions": ["1 gigabayt necha megabayt?", "Operatsion tizimga misol keltiring."],
    },
    "Internet va xavfsizlik asoslari": {
        "subject": "Informatika",
        "sections": [
            ("Internet nima?", "Internet — millionlab kompyuter va qurilmalarni bog'lovchi jahon tarmog'i. 1969-yilda ARPANET sifatida boshlangan.\n\nInternet xizmatlari:\n• Veb-saytlar (WWW)\n• Elektron pochta (e-mail)\n• Ijtimoiy tarmoqlar\n• Onlayn video va musiqa\n• Onlayn do'konlar"),
            ("Brauzer va veb-sayt", "Brauzer — internetdagi sahifalarni ko'rsatuvchi dastur.\nMisollar: Google Chrome, Firefox, Safari, Edge\n\nURL (manzil): https://www.google.com\n• https — xavfsiz protokol\n• www — jahon to'ri\n• google.com — domen nomi\n\nQidiruv tizimlari: Google, Yandex, Bing"),
            ("Internet xavfsizligi", "Xavfsiz internet qoidalari:\n• Parolni hech kimga bermang\n• Noma'lum havolalarga kirмang\n• Shaxsiy ma'lumotlarni internet'ga joylashtirishmang\n• Antivirus dastur o'rnating\n• Ijtimoiy tarmoqlarda begonalarga ishonmang\n\nFishing — soxta saytlar orqali ma'lumot o'g'irlash."),
        ],
        "terms": [("Internet", "Jahon kompyuter tarmog'i"), ("Brauzer", "Veb-sahifalarni ko'rsatuvchi dastur"), ("Fishing", "Soxta sayt orqali ma'lumot o'g'irlash")],
        "questions": ["Brauzerga misol keltiring.", "Xavfsiz internet parol qoidasi nima?"],
    },
    "Qog'oz va loy ishlari": {
        "subject": "Texnologiya",
        "sections": [
            ("Qog'oz bilan ishlash", "Qog'oz — eng ko'p ishlatiladigan xom ashyo. Qog'oz ishlari turlari:\n\n• Origami — qog'ozdan shakllar yasash (Yaponiya san'ati)\n• Applikatsiya — qog'oz parcha yapishtirib rasm yaratish\n• Quticha va kaftaklar yasash\n• Qog'oz mozaikasi\n\nQog'oz o'lchamlari: A4 (210×297mm), A3, A5"),
            ("Origami asoslari", "Origami — qog'ozni qirqmasdan, yopishmasdan faqat katlash orqali shakllar yasash.\n\nAsosiy katlash usullari:\n• To'g'ri katlash (valley fold)\n• Orqa katlash (mountain fold)\n\nOddiy origami shakllar: qayiq, kema, samolyot, qushcha, gul\n\nOrigami xotira, zehn va qo'l mahoratini rivojlantiradi."),
            ("Loy bilan ishlash", "Loy — tabiiy materiаl. Suv qo'shib yumshatiladi va istalgan shaklga kiritiladi.\n\nLoy ishlari bosqichlari:\n1. Loyni yumshatiш\n2. Shaklga keltirish\n3. Quritish yoki pishirish\n\nKulolchilik — loydan idish-tovoq yasash san'ati. O'zbekistonda Rishton, G'ijduvon kulolchiligi mashhur."),
        ],
        "terms": [("Origami", "Qog'oz katlash san'ati"), ("Applikatsiya", "Qog'oz parchalarini yapishtirib rasm yaratish"), ("Kulolchilik", "Loydan idish yasash san'ati")],
        "questions": ["Origami qaysi mamlakat san'ati?", "Loy bilan ishlash bosqichlarini ayting."],
    },
    "Mato va kashtachilik": {
        "subject": "Texnologiya",
        "sections": [
            ("Mato turlari", "Mato — to'qilgan material. Mato turlari:\n\n• Tabiiy matolar: paxta, jun, ipak, zig'ir\n• Sun'iy matolar: naylon, polyester, kapron\n\nPaxta matosi — о'zbekistonda eng ko'p ishlatiladigan. Ipak — qimmatbaho tabiiy mato, ipak qurtidan olinadi.\n\nMatolarning xususiyatlari: yumshoqlik, mustahkamlik, nafas o'tkazish."),
            ("Kashtachilik", "Kashtachilik — mato yoki boshqa materialga igna va ip bilan naqsh tushirish san'ati.\n\nO'zbekiston kashtachiligi — jahon miqyosida mashhur. Buxoro, Nurata, Shaxrisabz kashtachiligi alohida uslubga ega.\n\nKashtachilik asboblari: igna, ip (atlas, paxta, jun), dastgoh (gerchi).\n\nNaqsh turlari: gul, bargcha, geometrik shakllar, arxitektura elementlari."),
            ("Tikuv asoslari", "Tikuv — matoni bir-biriga ulash jarayoni. Tikuv usullari:\n\n• Qo'l bilan tikish — igna va ip bilan\n• Mashina bilan tikish — tikuv mashinasi orqali\n\nAsosiy tikuv turlari: to'g'ri chok, zanjir chok, naqsh chok.\n\nO'lchov olish, buyumni kesish va tikish — asosiy tikuvchilik bosqichlari."),
        ],
        "terms": [("Kastashchilik", "Mato va matoga naqsh tushirish san'ati"), ("Paxta", "Tabiiy o'simlik tolali mato"), ("Dastgoh", "Kashtachilik uchun maxsus ramka")],
        "questions": ["Ipak qayerdan olinadi?", "O'zbekistonning qaysi shaharlari kashtachiligi bilan mashhur?"],
    },
    "Ranglar va geometrik shakllar": {
        "subject": "Tasviriy san'at",
        "sections": [
            ("Ranglar nazariyasi", "Ranglar uch guruhga bo'linadi:\n\n1. Asosiy ranglar: qizil, sariq, ko'k\n2. Qo'shimcha ranglar: yashil (sariq+ko'k), to'q sariq (qizil+sariq), binafsha (qizil+ko'k)\n3. Neytral ranglar: oq, qora, kulrang\n\nIssiq ranglar: qizil, sariq, to'q sariq — haroratni his ettiradi.\nSovuq ranglar: ko'k, yashil, binafsha — salqinlikni his ettiradi."),
            ("Geometrik shakllar", "Geometrik shakllar — aniq o'lchamli matematik shakllar:\n\n2D shakllar:\n• Doira — markaz va radius\n• Kvadrat — teng 4 tomonli to'rtburchak\n• Uchburchak — uch burchakli shakl\n• To'rtburchak — to'rtta to'g'ri burchakli\n\n3D shakllar:\n• Kub, sharsimon, silindr, konus, piramida"),
            ("Naqsh va bezak", "Naqsh — takrorlanuvchi geometrik yoki o'simlik motivlaridan iborat bezak.\n\nO'zbek naqshlari turlari:\n• Islimiy — o'simlik motivlari, barglar, gullар\n• Girih — geometrik shakllar\n• Kundal — oltin va ranglar bilan bezash\n\nNaqshlar arxitektura, to'qimachilik, kulolchilikda keng qo'llaniladi."),
        ],
        "terms": [("Asosiy ranglar", "Qizil, sariq, ko'k — aralashtirish mumkin bo'lmagan ranglar"), ("Geometrik shakl", "Aniq o'lchamli matematik shakl"), ("Naqsh", "Takrorlanuvchi bezak elementi")],
        "questions": ["Sariq + ko'k = qaysi rang?", "Asosiy ranglarni ayting."],
    },
    "Natyurmort va peyzaj asoslari": {
        "subject": "Tasviriy san'at",
        "sections": [
            ("Natyurmort", "Natyurmort — harakatsiz narsalar (meva, sabzavot, idish, kitob va boshqalar) tasvirini chizish.\n\nNatyurmort yasash bosqichlari:\n1. Kompozitsiyani tanlash — narsalarni joylashtirish\n2. Umumiy shaklni chizish\n3. Tafsilotlarni chizish\n4. Soya va yorug'likni ko'rsatish\n5. Rangga bo'yash"),
            ("Peyzaj", "Peyzaj — tabiat ko'rinishini tasvirlovchi rasm turi. Peyzaj elementlari:\n• Osmоn va bulutlar\n• Tog', tepalik\n• Daryo, ko'l, dengiz\n• Daraxt va o'simliklar\n• Uy va yo'llar\n\nLiniyali perspektiva — uzoqda narsalar kichikroq ko'rinadi."),
            ("Rasm chizish texnikasi", "Asosiy rasm chizish texniкalari:\n\n• Qalamcha (karandash) — qoralamalar va eskiz\n• Akvarell — suvli bo'yoq, shaffof\n• Gouash — qalin, qoplaydigan bo'yoq\n• Moy bo'yoq — boy faktura, kuchli rang\n• Pastell — yumshoq, rang beradigan qalam\n\nHar bir texnikaning o'ziga xos xususiyatlari bor."),
        ],
        "terms": [("Natyurmort", "Harakatsiz narsalar tasviri"), ("Peyzaj", "Tabiat manzarasi tasviri"), ("Akvarell", "Suvli shaffof bo'yoq")],
        "questions": ["Natyurmort nima?", "Peyzajda nima tasvirlanadi?"],
    },
    "Algebraik ifodalar va tenglamalar": {
        "subject": "Algebra",
        "sections": [
            ("Algebraik ifodalar", "Algebra — matematik fan. Algebra so'zi Al-Xorazmiyning \"Al-jabr\" asaridan kelib chiqqan.\n\nAlgebraik ifoda — son, harf va arifmetik belgilardan tashkil topgan yozuv.\nMisol: 2x + 3, 5a - b, x² + y²\n\nKoeffitsient — harfning oldidagi son. 3x da koeffitsient = 3."),
            ("Tenglamalar", "Tenglama — ikkita ifodaning tengligini ko'rsatuvchi yozuv.\nMisol: x + 5 = 12\n\nTenglamani yechish — x ning qiymatini topish.\nx + 5 = 12 → x = 12 - 5 = 7\n\nTenglama ildizi — tenglamani qanoatlantiradigan son.\nTekshirish: 7 + 5 = 12 ✓"),
            ("Tengsizliklar", "Tengsizlik — ikkita ifodaning tengma-tengmasligini ko'rsatadi.\n• a > b — a katta b dan\n• a < b — a kichik b dan\n• a ≥ b — a katta yoki teng b ga\n• a ≤ b — a kichik yoki teng b ga\n\nMisol: x + 3 > 7 → x > 4"),
        ],
        "terms": [("Tenglama", "Ikkita ifodaning tengligini ko'rsatuvchi yozuv"), ("Ildiz", "Tenglamani qanoatlantiradigan son"), ("Koeffitsient", "Harfning oldidagi son")],
        "questions": ["x + 8 = 15 ni yeching.", "Algebra so'zi qayerdan kelib chiqqan?"],
    },
    "Manfiy sonlar va modullari": {
        "subject": "Algebra",
        "sections": [
            ("Manfiy sonlar", "Son o'qi — sonlarni chiziqda tasvirlash. Noldan chapda manfiy sonlar, o'ngda musbat sonlar.\n\n...-3, -2, -1, 0, 1, 2, 3...\n\nManfiy son — noldan kichik son. Masalan, -5, -0.7, -100.\n\nManfiy sonlar amaliyotda: sovuqda harorat (-10°C), dengiz sathidan past (-400 m)."),
            ("Manfiy sonlarda amallar", "Qo'shish:\n(+3) + (+5) = +8\n(-3) + (-5) = -8\n(+3) + (-5) = -2\n\nKo'paytirish:\n(+) × (+) = (+)\n(-) × (-) = (+)\n(+) × (-) = (-)\n\nMisol: (-4) × (-3) = +12\n(-4) × (+3) = -12"),
            ("Absolyut qiymat (modul)", "Modulь — sonning noldan masofasi.\n\n|5| = 5\n|-5| = 5\n|0| = 0\n\nModul har doim musbat yoki nol bo'ladi.\n\n|-7| + |3| = 7 + 3 = 10\n\nModul belgisi: |x| — ikki vertikal chiziq ichida."),
        ],
        "terms": [("Manfiy son", "Noldan kichik son"), ("Modul", "Sonning noldan masofasi — absolyut qiymat"), ("Son o'qi", "Sonlarni tasvirlovchi chiziq")],
        "questions": ["(-5) × (-3) = ?", "|-8| = ?"],
    },
    "O'zbek tili va imlo qoidalari": {
        "subject": "O'zbek tili",
        "sections": [
            ("O'zbek tili tarixi", "O'zbek tili — turkiy tillar oilasiga mansub. 30 milliondan ortiq kishi gaplashadi.\n\nYozuv tarixi:\n• Qadimgi turkiy yozuv — runik yozuv\n• Arab yozuvi — VIII–XX asr\n• Lotin yozuvi — 1929–1940, va 1995-yildan qayta\n• Kirill yozuvi — 1940–1995\n\nHozir rasmiy yozuv — lotin yozuvi."),
            ("Imlo qoidalari", "Imlo — to'g'ri yozish qoidalari.\n\nAsosiy qoidalar:\n• Katta harf: gap boshi, ism, joy nomlari\n• Qo'shib yozish: birga ma'no anglatgan so'zlar (bugun, shuning uchun)\n• Ajratib yozish: mustaqil ma'noli so'zlar\n• Tire: qarama-qarshi tushunchalar (qora-oq, issiq-sovuq)\n• Apоstrof: tovush tushishi (o'zbek, to'rt)"),
            ("Punktuatsiya", "Tinish belgilari — gapni to'g'ri o'qishga yordam beradi:\n\n• Nuqta (.) — darak gapning oxiri\n• So'roq belgisi (?) — so'roq gapning oxiri\n• Undov belgisi (!) — undov gapning oxiri\n• Vergul (,) — sanash, qo'shimcha ma'lumot\n• Ikki nuqta (:) — izohlash\n• Tire (—) — to'xtam, qarama-qarshilik"),
        ],
        "terms": [("Imlo", "To'g'ri yozish qoidalari"), ("Punktuatsiya", "Tinish belgilari tizimi"), ("Apostrof", "O'zbek yozuvida tovush tushishini ko'rsatuvchi belgi")],
        "questions": ["O'zbek tili qaysi til oilasiga kiradi?", "Katta harf qachon yoziladi?"],
    },
    "So'z yasalishi va sinonimlar": {
        "subject": "O'zbek tili",
        "sections": [
            ("So'z yasalishi", "Yangi so'zlar quyidagi yo'llar bilan yasaladi:\n\n1. Qo'shimcha usuli: kitob → kitobxon, ishla → ishlovchi\n2. Qo'shish usuli: ota-ona, qora-qo'ng'iz\n3. Qisqartirish: O'zbekiston → O'zbek, texnika → tex\n\nSo'z yasovchi qo'shimchalar:\n-chi (kasb): ishchi, dehqonchi\n-lik (xususiyat): kattаlik, yaxshilik\n-dosh (sheriklik): vatandosh, sinfдош"),
            ("Sinonimlar", "Sinonimlar — bir xil yoki yaqin ma'noli so'zlar.\n\nMisollar:\n• Katta — ulkan — yirik — bahaybat\n• Yaxshi — ajoyib — zo'r — a'lo\n• Tez — shoshqin — jadal — shiddat\n\nSinonimlar nutqni boyitadi va takrorni kamaytiradi. Sinonimik qator — bir ma'noni bildiruvchi so'zlar qatori."),
            ("Antonimlar va omonimlar", "Antonimlar — qarama-qarshi ma'noli so'zlar.\nMisol: baland–past, issiq–sovuq, katta–kichik, yaxshi–yomon\n\nOmonimlar — bir xil yoziladigan, lekin boshqa ma'noli so'zlar.\nMisol:\n• Oq — rang (oq ko'ylak) va oqmoq (suv oqadi)\n• Qoq — qoqmoq (teshik qoqmoq) va qoq (ozg'in)"),
        ],
        "terms": [("Sinonim", "Bir xil yoki yaqin ma'noli so'z"), ("Antonim", "Qarama-qarshi ma'noli so'z"), ("Omonim", "Bir xil yoziladigan boshqa ma'noli so'z")],
        "questions": ["\"Katta\" so'zining sinonimlarini ayting.", "Antonimga misol keltiring."],
    },
}


def slugify(title: str) -> str:
    s = title.lower()
    s = re.sub(r"[''`]", "", s)
    s = re.sub(r"[^a-z0-9а-яёo'ʻЀ-ӿ\s-]", "", s)
    s = re.sub(r"\s+", "_", s.strip())
    s = re.sub(r"[^a-z0-9_-]", "", s.encode("ascii", "ignore").decode())
    return s[:60] or "lesson"


def make_docx(title: str, subject: str, data: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 70, 127)

    # Subject
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Fan: {subject}  |  5-sinf")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    for heading, body in data.get("sections", []):
        h = doc.add_paragraph()
        hr = h.add_run(heading)
        hr.font.name = "Times New Roman"
        hr.font.size = Pt(13)
        hr.font.bold = True
        hr.font.color.rgb = RGBColor(0, 102, 51)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)

        for para in body.split("\n\n"):
            bp = doc.add_paragraph(para.strip())
            if bp.runs:
                br = bp.runs[0]
                br.font.name = "Times New Roman"
                br.font.size = Pt(12)
            bp.paragraph_format.space_after = Pt(6)
            bp.paragraph_format.line_spacing = Pt(18)

    # Terms
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tr = tp.add_run("Asosiy tushunchalar")
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(13)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0, 102, 51)

    for term, definition in data.get("terms", []):
        p = doc.add_paragraph()
        pr = p.add_run(f"{term}: ")
        pr.font.name = "Times New Roman"
        pr.font.size = Pt(12)
        pr.font.bold = True
        dr = p.add_run(definition)
        dr.font.name = "Times New Roman"
        dr.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(3)

    # Questions
    doc.add_paragraph()
    qp = doc.add_paragraph()
    qr = qp.add_run("Savol va topshiriqlar")
    qr.font.name = "Times New Roman"
    qr.font.size = Pt(13)
    qr.font.bold = True
    qr.font.color.rgb = RGBColor(0, 102, 51)

    for i, q in enumerate(data.get("questions", []), 1):
        p = doc.add_paragraph(f"{i}. {q}")
        if p.runs:
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(12)
        p.paragraph_format.space_after = Pt(3)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("5-sinf o'quv materiali  |  2025–2026 o'quv yili")
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(150, 150, 150)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Map title → filename, write files
    mapping = {}
    saved = 0
    skipped = 0

    for title, data in LESSONS.items():
        fname = f"{slugify(title)}.docx"
        fpath = OUTPUT_DIR / fname
        content = make_docx(title, data["subject"], data)
        fpath.write_bytes(content)
        mapping[title] = f"/api/files/lessons/{fname}"
        saved += 1
        print(f"  OK  {fname}")

    print(f"\nSaved {saved} files to {OUTPUT_DIR}")
    print("\nMapping (title → pdf_url):")
    for t, u in mapping.items():
        print(f"  {t!r}: {u!r}")

    # Write mapping JSON for TS seed script
    import json
    map_path = Path(__file__).parent / "lesson_file_mapping.json"
    with open(map_path, "w", encoding="utf-8") as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)
    print(f"\nMapping saved: {map_path}")


if __name__ == "__main__":
    main()
