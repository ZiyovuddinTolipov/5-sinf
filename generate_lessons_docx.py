"""
Generate sample .docx lesson files for each topic in tests.json.
Requires: pip install python-docx
Output: lessons/ folder with one .docx per topic
"""

import json
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Install python-docx first: pip install python-docx")
    raise

LESSONS_CONTENT = {
    "1-Mavzu: Tarixga kirish": {
        "title": "Tarixga kirish",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Tarix fani nima?",
                "body": (
                    "Tarix — insoniyat o'tmishini o'rganuvchi fandir. U qadimgi davrlardan "
                    "bugungi kungacha bo'lgan voqealarni, odamlarning hayoti, madaniyati, "
                    "davlatlari va sivilizatsiyalarini o'rganadi.\n\n"
                    "\"Tarix\" so'zi yunoncha \"historia\" so'zidan olingan bo'lib, \"tadqiqot\" "
                    "yoki \"bilim\" degan ma'noni anglatadi. Tarix fanining otasi — qadimgi yunon "
                    "tarixchisi Gerodot hisoblanadi."
                ),
            },
            {
                "heading": "Tarixiy manbalar",
                "body": (
                    "Tarixchilar o'tmishni o'rganish uchun turli manbalardan foydalanadilar:\n\n"
                    "• Moddiy manbalar: qazishmalarda topilgan buyumlar, qurollar, binolar.\n"
                    "• Yozma manbalar: qo'lyozmalar, kitoblar, qonunlar, maktublar.\n"
                    "• Og'zaki manbalar: xalq og'zaki ijodi, dostonlar, rivoyatlar.\n"
                    "• Tasviriy manbalar: rasmlar, naqshlar, haykaltaroshlik asarlari."
                ),
            },
            {
                "heading": "Tarix yordamchi fanlari",
                "body": (
                    "Tarixni o'rganishda bir qator yordamchi fanlar ko'maklashadi:\n\n"
                    "• Arxeologiya — yerdan topilgan ashyolarni o'rganadi.\n"
                    "• Numizmatika — tangalar va medallarni o'rganadi.\n"
                    "• Genealogiya — nasl-nasab, oila daraxtlarini o'rganadi.\n"
                    "• Geraldika — gerblar va timsollarni o'rganadi.\n"
                    "• Paleografiya — qadimgi yozuvlarni o'rganadi."
                ),
            },
            {
                "heading": "Vaqt hisobi",
                "body": (
                    "Tarixda vaqt eramizdan avval (e.av.) va eramiz (e.) deb ikki qismga "
                    "bo'linadi. Eramizdan avvalgi yillar kamayib boradi: e.av. 100 yil — "
                    "e.av. 50 yil — e.av. 1 yil, keyin eramiz boshlanadi.\n\n"
                    "1 asr = 100 yil. 1 ming yillik = 10 asr = 1000 yil.\n\n"
                    "Misol: e.av. 3000-yil — bu bizdan 5000 yil oldin bo'lgan."
                ),
            },
        ],
        "key_terms": [
            ("Tarix", "Insoniyat o'tmishini o'rganuvchi fan"),
            ("Arxeologiya", "Qazishma ashyolarini o'rganuvchi fan"),
            ("Numizmatika", "Tangalarni o'rganuvchi fan"),
            ("Gerodot", "Tarix fanining otasi"),
            ("Asr", "100 yillik davr"),
        ],
        "questions": [
            "Tarix fani nimani o'rganadi?",
            "Tarixning otasi kim hisoblanadi?",
            "Numizmatika qanday fandir?",
            "Bir asr necha yilga teng?",
            "Eramizdan avvalgi yillar qanday sanaladi?",
        ],
    },
    "2-Mavzu: Ibtidoiy jamoa tuzumi": {
        "title": "Ibtidoiy jamoa tuzumi",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Ibtidoiy odamlarning paydo bo'lishi",
                "body": (
                    "Eng qadimgi odamlar Afrikada paydo bo'lgan. Ular hayvonlarga o'xshab "
                    "to'rt oyoqda yurganlar, keyinchalik asta-sekin tik yura boshlaganlar. "
                    "Ibtidoiy odamlar dastlab g'or va daraxtlarda yashaganlar.\n\n"
                    "O'zbekiston hududida topilgan eng qadimgi odam manzili — Farg'ona "
                    "vodiysidagi Selungur g'oridir (taxminan 1 million yil oldin)."
                ),
            },
            {
                "heading": "Mehnat qurollari",
                "body": (
                    "Ibtidoiy odamlar mehnat qurollarini toshdan yasaganlar. Tarix fanida "
                    "tosh davri uch bosqichga bo'linadi:\n\n"
                    "• Paleolit (Qadimgi tosh davri) — eng qo'pol tosh qurollar.\n"
                    "• Mezolit (O'rta tosh davri) — kamalak va o'q ixtiro qilindi.\n"
                    "• Neolit (Yangi tosh davri) — kulolchilik, to'qimachilik paydo bo'ldi.\n"
                    "• Eneolit (Mis-tosh davri) — mis metalidan foydalanish boshlandi."
                ),
            },
            {
                "heading": "Ibtidoiy jamoa tuzumi",
                "body": (
                    "Ibtidoiy odamlar guruh-guruh bo'lib yashaganlar. Bu guruhlar \"ibtidoiy "
                    "to'da\" deb atalgan. Keyinchalik qon-qarindoshlik asosida \"urug' jamoasi\" "
                    "shakllandi.\n\n"
                    "Jamoadagi barcha ishlar birgalikda bajarilgan: ov, baliq tutish, "
                    "meva-sabzavot yig'ish. Boshqaruvni oqsoqollar amalga oshirgan."
                ),
            },
            {
                "heading": "O'zbekistondagi ibtidoiy manzillar",
                "body": (
                    "• Selungur g'ori (Farg'ona) — 1 million yil oldin.\n"
                    "• Teshiktosh g'ori (Surxondaryo) — 70-100 ming yil oldin. "
                    "Bu yerda neandertal bolaning qoldiq suyaklari topilgan.\n"
                    "• Machay g'ori (Surxondaryo) — ibtidoiy odamlar yashaган.\n"
                    "• Zarautsoy (Surxondaryo) — qoya rasmlari topilgan."
                ),
            },
        ],
        "key_terms": [
            ("Paleolit", "Qadimgi tosh davri"),
            ("Mezolit", "O'rta tosh davri"),
            ("Neolit", "Yangi tosh davri"),
            ("Eneolit", "Mis-tosh davri"),
            ("Selungur", "O'zbekistondagi eng qadimgi odam manzili"),
        ],
        "questions": [
            "Ibtidoiy odamlarning ilk mehnat quroli nimadan yasalgan?",
            "Teshiktosh g'ori qayerda joylashgan?",
            "Kamalak va o'q qaysi davrda ixtiro qilingan?",
            "Urug' jamoasi qanday asosda birlashgan?",
            "Mis-tosh davri qanday nomlanadi?",
        ],
    },
    "3-Mavzu: Qadimgi Misr": {
        "title": "Qadimgi Misr",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Qadimgi Misr sivilizatsiyasi",
                "body": (
                    "Qadimgi Misr sivilizatsiyasi Nil daryosi bo'yida paydo bo'lgan. Nil har "
                    "yili toshib, qirg'oqlarga hosildor loy qoldirgan. Shu tufayli Misr "
                    "qadimdan dehqonchilik mamlakati bo'lgan.\n\n"
                    "Eramizdan avvalgi 3000-yilda Yuqori va Quyi Misr birlashib yagona "
                    "davlat tashkil etgan. Misr podshohlariga fir'avn deyilgan."
                ),
            },
            {
                "heading": "Ehromlar va Sfinks",
                "body": (
                    "Qadimgi Misrning eng mashhur inshootlari — ehromlar (piramidalar). "
                    "Eng katta ehrom — Xeops ehromining balandligi 147 metrga yetgan. "
                    "U taxminan 2,5 million tosh bloklardan qurilgan.\n\n"
                    "Giza platosida uchta katta ehrom bor: Xeops, Xafra va Mikerin. "
                    "Yaqinida mashhur Sfinks haykali joylashgan — arslon tanali, odam boshli haykal."
                ),
            },
            {
                "heading": "Misr yozuvi va fani",
                "body": (
                    "Misrliklar iyeroglif yozuvini ixtiro qilganlar. Iyeroglif — rasm-belgilardan "
                    "iborat yozuv. Ular papirus o'simligidan qog'oz tayyorlashgan.\n\n"
                    "Misrliklar astronomiya, matematika va tibbiyotda katta yutuqlarga erishganlar. "
                    "Ular yil davomiyligini 365 kun deb aniq hisoblashgan."
                ),
            },
            {
                "heading": "Misr dini va mumiyolash",
                "body": (
                    "Qadimgi Misrliklar ko'p xudoylarga sig'inganlar. Eng muhim xudolar: "
                    "Ra (quyosh xudosi), Osiris (o'limdan keyingi hayot xudosi), Isis (ona xudosi).\n\n"
                    "Misrliklar o'likni mumiyolab saqlashgan. Ular ruh o'limdan keyin ham "
                    "yashamoqchi bo'lsa, tana saqlanishi kerak deb ishonganlar."
                ),
            },
        ],
        "key_terms": [
            ("Fir'avn", "Misr podshohining unvoni"),
            ("Iyeroglif", "Qadimgi Misr yozuvi"),
            ("Papirus", "Misr o'simligidan tayyorlangan qog'oz"),
            ("Sfinks", "Arslon tanali, odam boshli haykal"),
            ("Mumiya", "Maxsus usulda saqlangan o'lik"),
        ],
        "questions": [
            "Misr qaysi daryo bo'yida joylashgan?",
            "Eng katta ehrom nomi nima?",
            "Misr yozuvi qanday nomlanadi?",
            "Misr qog'ozi nima deb ataladi?",
            "Sfinks nima?",
        ],
    },
    "4-Mavzu: Ikki daryo oralig'i": {
        "title": "Ikki daryo oralig'i — Mesopotamiya",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Mesopotamiya sivilizatsiyasi",
                "body": (
                    "Mesopotamiya (\"ikki daryo oralig'i\") — Dajla (Tigr) va Furot (Evfrat) "
                    "daryolari orasidagi hudud. Bugungi Iroq davlatining asosiy qismi shu "
                    "hududda joylashgan.\n\n"
                    "Mesopotamiyada eng qadimgi sivilizatsiyalardan biri — Shumer sivilizatsiyasi "
                    "paydo bo'lgan. Shumerlar eramizdan avvalgi IV ming yillikda yashashgan."
                ),
            },
            {
                "heading": "Shumerlarning kashfiyotlari",
                "body": (
                    "Shumerlar insoniyatga juda ko'p muhim ixtirolar berganlar:\n\n"
                    "• Mixxat yozuvi — loyga o'tkir buyum bilan yozilar edi.\n"
                    "• G'ildirak — taxminan e.av. 3500-yil ixtiro qilingan.\n"
                    "• Zikkurat — baland ibodatxona minorasi.\n"
                    "• Mirzalik (yozuv maktablari) — dunyodagi birinchi maktablar."
                ),
            },
            {
                "heading": "Bobil va Hammurapi qonunlari",
                "body": (
                    "Mesopotamiyaning eng qudratli shahri — Bobil. Bobil shohi Hammurapi "
                    "(e.av. 1792–1750 yillar) dunyadagi birinchi yozma qonunlar to'plamini "
                    "tuzgan. Bu qonunlar tosh ustunlarga o'yib yozilgan.\n\n"
                    "Hammurapi qonunlarida 282 ta modda bor. Asosiy tamoyil: \"Ko'zga ko'z, "
                    "tishga tish.\" Qonunlar savdo, nikoh, mulk, jinoyat masalalarini tartibga solgan."
                ),
            },
            {
                "heading": "Ashshurbanipal kutubxonasi",
                "body": (
                    "Ossuriya shohi Ashshurbanipal (e.av. 668–627) dunyodagi birinchi "
                    "kutubxonani tashkil etgan. Nineviya shahrida joylashgan bu kutubxonada "
                    "30 000 dan ortiq loy lavhalar saqlangan.\n\n"
                    "\"Gilgamish haqida doston\" — insoniyat tarixidagi eng qadimgi yozma "
                    "adabiy asar. U shu kutubxonadan topilgan."
                ),
            },
        ],
        "key_terms": [
            ("Mesopotamiya", "Dajla va Furot orasidagi hudud"),
            ("Mixxat", "Shumerlarning loy lavhaga yoziladigan yozuvi"),
            ("Zikkurat", "Shumer ibodatxonasi minorasi"),
            ("Hammurapi", "Bobil shohi, birinchi qonunlar muallifi"),
            ("Ashshurbanipal", "Birinchi kutubxonani tashkil etgan shoh"),
        ],
        "questions": [
            "Ikki daryo oralig'i qaysi daryolar orasida joylashgan?",
            "Shumerlar nimani ixtiro qilganlar?",
            "Hammurapi qonunlarida nechta modda bor?",
            "Dunyodagi birinchi kutubxona kimga tegishli?",
            "\"Gilgamish haqida doston\" qayerdan topilgan?",
        ],
    },
    "5-Mavzu: Qadimgi Hindiston va Xitoy": {
        "title": "Qadimgi Hindiston va Xitoy",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Qadimgi Hindiston",
                "body": (
                    "Hindiston sivilizatsiyasi Hind va Gang daryolari vodiylarida rivojlangan. "
                    "Mohenjodaro va Xarappa — Hind vodiysi sivilizatsiyasining eng yirik shaharlari "
                    "(e.av. 2500-1500 yillar).\n\n"
                    "Hindistonda kasta tizimi mavjud edi. Jamiyat to'rt kastaga bo'lingan: "
                    "brahmanlar (kahinlar), kshatriylar (jangchilar), vayshyalar (savdogarlar, "
                    "dehqonlar), shudralar (xizmatkorlar)."
                ),
            },
            {
                "heading": "Hindiston madaniyati",
                "body": (
                    "Hindlar matematikada katta yutuqlarga erishdilar. Ular \"nol\" raqamini "
                    "ixtiro qilganlar. Bugungi biz foydalanadigan raqamlar (arab raqamlari) "
                    "aslida hind raqamlaridir.\n\n"
                    "Budda dini Hindistonda e.av. V asrda paydo bo'lgan. Budda — Siddxartha "
                    "Gautama laqabi. Budda ta'limoti: azob-uqubatdan qutulish yo'li — "
                    "tiyilish va meditatsiya."
                ),
            },
            {
                "heading": "Qadimgi Xitoy sivilizatsiyasi",
                "body": (
                    "Xitoy sivilizatsiyasi Xuanxe (\"Sariq daryo\") va Yangtsze daryolari "
                    "bo'yida rivojlangan. E.av. III ming yillikda Xitoyda yozuv paydo bo'lgan.\n\n"
                    "Xitoyliklar ko'plab muhim ixtirolarni dunyoga berishgan: qog'oz, bosma "
                    "mashina, porox, kompas, ipak, chinni (farfor)."
                ),
            },
            {
                "heading": "Xitoy devori",
                "body": (
                    "Buyuk Xitoy devori — e.av. III asrdan qurilgan. Uzunligi 21 000 km dan "
                    "oshiq. Maqsadi — shimoldan keluvchi ko'chmanchi qabilalardan himoya qilish.\n\n"
                    "Konfutsiy (e.av. 551–479) — Xitoyning buyuk faylasufi. Uning ta'limoti "
                    "(konfutsiylik) oila, ota-onaga hurmat, axloq va ta'lim-tarbiyaga asoslangan."
                ),
            },
        ],
        "key_terms": [
            ("Mohenjodaro", "Hind vodiysi sivilizatsiyasining yirik shahri"),
            ("Kasta", "Hindiston jamiyatidagi tabaqa tizimi"),
            ("Budda", "Siddxartha Gautama, budda dinining asoschisi"),
            ("Konfutsiy", "Qadimgi Xitoy faylasufi"),
            ("Buyuk Xitoy devori", "Dunyo mo'jizalaridan biri, uzunligi 21 000+ km"),
        ],
        "questions": [
            "Hind vodiysi sivilizatsiyasining yirik shaharlari qaysilar?",
            "Hindlar matematikada qanday muhim kashfiyot qilganlar?",
            "Xitoyliklar nimalarni ixtiro qilganlar?",
            "Buyuk Xitoy devori qanday maqsadda qurilgan?",
            "Konfutsiy kimdir?",
        ],
    },
    "6-Mavzu: O'rta Osiyo va Zardushtiylik": {
        "title": "O'rta Osiyo va Zardushtiylik",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Qadimgi O'rta Osiyo sivilizatsiyasi",
                "body": (
                    "O'rta Osiyo — Amudaryo va Sirdaryo orasidagi ulkan hudud. Bu hududda "
                    "qadimdan ko'plab davlatlar mavjud bo'lgan: Baqtriya, Sug'd, Xorazm, "
                    "Parfiya.\n\n"
                    "O'rta Osiyo \"Ipak yo'li\"ning markazida joylashgan. Bu savdo yo'li "
                    "Xitoyni Yaqin Sharq va O'rta yer dengizi mintaqasiga bog'lagan."
                ),
            },
            {
                "heading": "Zardushtiylik dini",
                "body": (
                    "Zardushtiylik — dunyodagi eng qadimgi dinlardan biri. Uni Zardusht "
                    "(Zaratushtra) e.av. VII–VI asrlarda targ'ib qilgan.\n\n"
                    "Zardushtiylikning muqaddas kitobi — Avesto. Asosiy tamoyil: yaxshilik "
                    "va yomonlik o'rtasidagi kurash. Yaxshi fikr, yaxshi so'z, yaxshi ish "
                    "(Humata, Huxta, Huvarshta) — zardushtiylikning uch asosi.\n\n"
                    "Zardushtiylikda olov muqaddas hisoblanadi. Ibodat joylari — otashkadalar."
                ),
            },
            {
                "heading": "Xorazm va Sug'd davlatlari",
                "body": (
                    "Xorazm — Amudaryo quyi oqimida joylashgan qadimgi davlat. Xorazmda "
                    "dehqonchilik va hunarmandchilik rivojlangan. Xorazm o'zining alohida "
                    "yozuviga ega bo'lgan.\n\n"
                    "Sug'd davlati — Zarafshon daryosi bo'yida joylashgan. Sug'diyona "
                    "savdogarlari Ipak yo'lida faol ishtirok etganlar. Samarqand va Buxoro "
                    "Sug'dning yirik shaharlari bo'lgan."
                ),
            },
            {
                "heading": "Moddiy madaniyat",
                "body": (
                    "O'rta Osiyo aholisi qadimdan sug'orma dehqonchilik bilan shug'ullangan. "
                    "Suv inshootlari — kanallar va arxlar qurilgan.\n\n"
                    "Yassi tepa, Afrosiyob (Samarqand), Topraqqal'a (Xorazm) — arxeologik "
                    "tadqiqotlar o'tkazilgan qadimgi shahar qoldiqlari. Bu yerlardan sopol "
                    "buyumlar, tanga pullar, haykalchalar topilgan."
                ),
            },
        ],
        "key_terms": [
            ("Zardushtiylik", "Zardusht ta'limotiga asoslangan qadimgi din"),
            ("Avesto", "Zardushtiylikning muqaddas kitobi"),
            ("Ipak yo'li", "Xitoyni G'arb bilan bog'lagan savdo yo'li"),
            ("Sug'd", "Zarafshon bo'yidagi qadimgi davlat"),
            ("Xorazm", "Amudaryo quyi oqimidagi qadimgi davlat"),
        ],
        "questions": [
            "Zardushtiylikning muqaddas kitobi nima?",
            "Zardusht qachon yashagan?",
            "Zardushtiylikning uch asosi qaysilar?",
            "Ipak yo'li qanday yo'nalishda o'tgan?",
            "Sug'd davlati qayerda joylashgan?",
        ],
    },
    "7-Mavzu: Qadimgi Yunoniston": {
        "title": "Qadimgi Yunoniston",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Yunon sivilizatsiyasi",
                "body": (
                    "Qadimgi Yunoniston Bolqon yarim orolida joylashgan. Dengizga yaqinligi "
                    "tufayli yunonlar dengizchilik va savdoni rivojlantirganlar. Yunonlar "
                    "butun O'rta Yer dengizi bo'ylab koloniyalar tashkil etganlar.\n\n"
                    "Yunoniston ko'plab shahar-davlatlarga (polislar) bo'lingan edi. "
                    "Eng muhim polislar: Afina va Sparta."
                ),
            },
            {
                "heading": "Afina demokratiyasi",
                "body": (
                    "Afina — demokratiyaning vatani. E.av. V asrda Perikl boshchiligida "
                    "Afinada erkin fuqarolar davlat ishlarini hal qilganlar.\n\n"
                    "\"Demokratiya\" so'zi yunoncha \"demos\" (xalq) + \"kratos\" (hokimiyat) "
                    "so'zlaridan hosil bo'lgan. Xalq majlisi — Ekklesia — eng oliy "
                    "hokimiyat organi hisoblangan."
                ),
            },
            {
                "heading": "Yunon fani va madaniyati",
                "body": (
                    "Qadimgi Yunoniston fan va madaniyatda ulkan muvaffaqiyatlarga erishgan:\n\n"
                    "• Faylasuflar: Sokrat, Platon, Arastu.\n"
                    "• Tarixchilar: Gerodot (tarix otasi), Fukidid.\n"
                    "• Matematik: Pifagor (Pifagor teoremasi), Evklid (geometriya).\n"
                    "• Tibbiyot: Gippokrat (tibbiyot otasi).\n"
                    "• Dramaturgia: Sofokl, Esxil, Aristofan."
                ),
            },
            {
                "heading": "Olimpiya o'yinlari",
                "body": (
                    "Qadimgi yunon olimpiya o'yinlari e.av. 776-yilda boshlanib, har 4 yilda "
                    "bir marta o'tkazilgan. O'yinlar Zeus xudosiga bag'ishlangan.\n\n"
                    "Musobaqalar turlari: yugurish, kurash, ot poygasi, disk uloqtirish, "
                    "nayzapardoz. O'yin davomida urushlar to'xtatilgan — \"olimpiya tinchligi.\""
                ),
            },
        ],
        "key_terms": [
            ("Polis", "Yunoniston shahar-davlati"),
            ("Demokratiya", "Xalq hokimiyati — Afinada paydo bo'lgan boshqaruv shakli"),
            ("Olimpiya o'yinlari", "Har 4 yilda bir marta o'tkazilgan sport musobaqalari"),
            ("Arastu", "Qadimgi yunon faylasufi va olimi"),
            ("Gippokrat", "Tibbiyot fanining otasi"),
        ],
        "questions": [
            "Demokratiya qayerda paydo bo'lgan?",
            "\"Demokratiya\" so'zi nima ma'noni anglatadi?",
            "Olimpiya o'yinlari qachon boshlangan?",
            "Tibbiyot fanining otasi kim?",
            "Tarix fanining otasi kim?",
        ],
    },
    "8-Mavzu: Aleksandr Makedonskiy": {
        "title": "Aleksandr Makedonskiy",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Aleksandrning hayoti",
                "body": (
                    "Aleksandr Makedonskiy e.av. 356-yilda tug'ilgan. Uning otasi — Makedoniya "
                    "shohi Filip II. Aleksandr mashhur faylasuf Arastu shogirdi bo'lgan.\n\n"
                    "E.av. 336-yilda, 20 yoshida Aleksandr Makedoniya taxtiga o'tirgan. "
                    "Katta armiya bilan u Fors imperiyasiga qarshi yurish boshlagan."
                ),
            },
            {
                "heading": "Aleksandrning yurishlari",
                "body": (
                    "E.av. 334-yilda Aleksandr Osiyoga o'tdi. U quyidagi hududlarni zabt etdi:\n\n"
                    "• Kichik Osiyo (hozirgi Turkiya)\n"
                    "• Suriya va Falastin\n"
                    "• Misr — u yerda Iskandariya (Aleksandriya) shahrini asos soldi\n"
                    "• Fors imperiyasi (hozirgi Eron)\n"
                    "• Baqtriya va Sug'd (hozirgi O'zbekiston va Afg'oniston)\n"
                    "• Shimoliy Hindiston"
                ),
            },
            {
                "heading": "O'rta Osiyoda Aleksandr",
                "body": (
                    "E.av. 329-yilda Aleksandr O'rta Osiyoga kirib keldi. U Marakanda "
                    "(Samarqand) va So'g'diana hududini bosib oldi.\n\n"
                    "Mahalliy aholi Spitamen boshchiligida Aleksandrga qarshi qo'zg'olon "
                    "ko'tardi. Uch yillik kurash davomida Aleksandr O'rta Osiyoni qisman "
                    "nazorat ostiga oldi."
                ),
            },
            {
                "heading": "Aleksandrning merosi",
                "body": (
                    "Aleksandr e.av. 323-yilda 32 yoshida Bobilda vafot etdi. U asos solgan "
                    "imperiya uning o'limidan keyin uch qismga bo'lindi.\n\n"
                    "Aleksandrning merosi: yunon-sharq madaniyatining (ellinizm) tarqalishi. "
                    "Yunon tili, san'ati va falsafasi Sharq bilan qo'shilib yangi madaniyat "
                    "yaratdi. Iskandariya — o'sha davrning ilm-fan markazi bo'ldi."
                ),
            },
        ],
        "key_terms": [
            ("Makedoniya", "Yunon yarim orolining shimoliy qismi"),
            ("Ellinizm", "Yunon-sharq madaniyatining sintezi"),
            ("Iskandariya", "Aleksandr asos solgan Misrdagi shahar"),
            ("Spitamen", "O'rta Osiyoda Aleksandrga qarshi kurashgan qo'mondon"),
            ("Marakanda", "Samarqandning qadimgi nomi"),
        ],
        "questions": [
            "Aleksandr Makedonskiy qaysi faylasuf shogirdi bo'lgan?",
            "Aleksandr Misrda qanday shahar asos solgan?",
            "Aleksandr O'rta Osiyoga qachon kirib kelgan?",
            "Marakanda — bu qaysi shaharning qadimgi nomi?",
            "Aleksandr qanday yoshda vafot etgan?",
        ],
    },
    "9-Mavzu: Qadimgi Rim": {
        "title": "Qadimgi Rim",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Rim sivilizatsiyasi",
                "body": (
                    "Qadimgi Rim e.av. 753-yilda Italiya yarim orolida asos solingan. "
                    "Rim afsonasiga ko'ra, shaharni egizaklar Romul va Rem asos solgan. "
                    "Romul birinchi rim shohiga aylangan.\n\n"
                    "Rim dastlab shahar-davlat bo'lib, keyin ulkan imperiyaga aylandi. "
                    "O'z eng yuqori cho'qqisida Rim O'rta yer dengizini to'liq nazorat qilgan."
                ),
            },
            {
                "heading": "Rim respublikasi",
                "body": (
                    "E.av. 509-yilda shoh quvib chiqarildi va respublika o'rnatildi. "
                    "Hokimiyat ikki konsul va Senat qo'lida bo'lgan.\n\n"
                    "Rim huquqi — antik dunyoning eng rivojlangan huquq tizimi. Rim "
                    "qonunlari ko'plab zamonaviy huquq tizimlarining asosini tashkil etadi. "
                    "\"XII jadval qonunlari\" (e.av. 451-450) — birinchi yozma rim qonunlari."
                ),
            },
            {
                "heading": "Rim imperiyasi",
                "body": (
                    "E.av. 27-yilda Avgust Tsezar birinchi rim imperatori bo'ldi. Rim "
                    "imperiyasi o'z cho'qqisida Britaniyadan Mesopotamiyagacha cho'zilgan.\n\n"
                    "Rim inshootlari: yo'llar (barcha yo'llar Rimga olib boradi), "
                    "suv quvurlari (akveduklar), Kolizey (gladiatorlar jangi uchun), "
                    "Panteon (barcha xudolarga ibodatxona)."
                ),
            },
            {
                "heading": "Rim madaniyati va dini",
                "body": (
                    "Rimliklar yunon madaniyatini o'zlashtirdi va rivojlantirdi. Lotin tili "
                    "imperiyaning rasmiy tili bo'lib, keyinchalik roman tillarining asosiga aylandi.\n\n"
                    "Xristianlik I asrda Falastinda paydo bo'ldi. Dastlab ta'qib qilindi, "
                    "keyinchalik 313-yilda imperator Konstantin tomonidan rasmiy din deb tan olindi."
                ),
            },
        ],
        "key_terms": [
            ("Respublika", "Senat va ikkita konsul boshqargan davlat shakli"),
            ("Senator", "Rim Senati a'zosi"),
            ("Imperator", "Rim davlatining oliy hukmdori"),
            ("Gladiator", "Rimda jang ko'rsatgan qul"),
            ("Lotin tili", "Rim imperiyasining rasmiy tili"),
        ],
        "questions": [
            "Qadimgi Rim qachon asos solingan?",
            "\"XII jadval qonunlari\" nima?",
            "Birinchi rim imperatori kim?",
            "Kolizey nima uchun qurilgan?",
            "Xristianlik Rimda qachon rasmiy din deb tan olindi?",
        ],
    },
    "10-Mavzu: Dovon va Qang' davlatlari": {
        "title": "Dovon va Qang' davlatlari",
        "subtitle": "5-sinf. Tarix fani",
        "sections": [
            {
                "heading": "Dovon davlati",
                "body": (
                    "Dovon davlati — hozirgi Farg'ona vodiysida joylashgan qadimgi davlat. "
                    "Xitoy manbalarida \"Davan\" nomi bilan tilga olingan.\n\n"
                    "Dovon e.av. II–I asrlarda mavjud bo'lgan. Paytaxti — Ershi (hozirgi "
                    "Marhamat shahrchasi atrofida). Dovon aholisi asosan dehqonchilik va "
                    "chorvachilik bilan shug'ullangan."
                ),
            },
            {
                "heading": "Dovonning \"Samoviy otlari\"",
                "body": (
                    "Dovon o'zining mashhur \"samoviy otlari\" (arғumoq otlar) bilan "
                    "mashxur bo'lgan. Xitoy imperatori bu otlarni qo'lga kiritish uchun "
                    "harbiy yurish uyushtirgan (e.av. 104-102 yillar).\n\n"
                    "Xitoyliklar Dovonni zabt eta olmagan, lekin katta miqdorda baholi "
                    "otlarni sovg'a sifatida olganlar. Bu voqea O'rta Osiyo va Xitoy "
                    "o'rtasidagi dastlabki diplomatik aloqalardan biri."
                ),
            },
            {
                "heading": "Qang' davlati",
                "body": (
                    "Qang' davlati — Sirdaryo o'rta va quyi oqimlarida joylashgan. "
                    "Xitoy manbalari bu davlatni \"Kangju\" deb atagan.\n\n"
                    "Qang' e.av. II asrdan eramizning IV asrigacha mavjud bo'lgan. "
                    "Paytaxti — Bityan shahri. Qang' aholisi ko'chmanchi va yarim "
                    "ko'chmanchi hayot kechirishgan."
                ),
            },
            {
                "heading": "Ipak yo'li va savdo",
                "body": (
                    "Dovon va Qang' davlatlari Ipak yo'lining muhim nuqtalarida joylashgan. "
                    "Bu davlatlar orqali Xitoy, O'rta Osiyo va G'arbni bog'lovchi savdo "
                    "karvonlari o'tgan.\n\n"
                    "Savdo: Xitoydan ipak, ziravorlar; G'arbdan shisha buyumlar, "
                    "qimmatbaho toshlar; O'rta Osiyodan at, qo'y va qo'l hunarmandchiligi "
                    "mahsulotlari almashilgan."
                ),
            },
        ],
        "key_terms": [
            ("Dovon", "Hozirgi Farg'ona vodiysidagi qadimgi davlat"),
            ("Qang'", "Sirdaryo bo'yidagi qadimgi davlat"),
            ("Samoviy otlar", "Dovonning mashhur arğumoq otlari"),
            ("Ershi", "Dovon davlatining paytaxti"),
            ("Bityan", "Qang' davlatining paytaxti"),
        ],
        "questions": [
            "Dovon davlati qayerda joylashgan?",
            "Dovon nima bilan mashhur bo'lgan?",
            "Qang' davlati qaysi daryo bo'yida joylashgan?",
            "Ipak yo'li qanday savdoni ta'minlagan?",
            "Xitoy imperatori Dovonga nima uchun yurish uyushtirgan?",
        ],
    },
}


def set_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(0, 70, 127))
    elif level == 2:
        set_font(run, size=13, bold=True, color=(0, 102, 51))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text):
    for para_text in text.split("\n\n"):
        p = doc.add_paragraph(para_text.strip())
        run = p.runs[0] if p.runs else p.add_run(para_text.strip())
        set_font(run, size=12)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(18)


def add_key_terms(doc, terms):
    add_heading(doc, "Asosiy tushunchalar", level=2)
    for term, definition in terms:
        p = doc.add_paragraph()
        run_term = p.add_run(f"{term}: ")
        set_font(run_term, size=12, bold=True)
        run_def = p.add_run(definition)
        set_font(run_def, size=12)
        p.paragraph_format.space_after = Pt(4)


def add_questions(doc, questions):
    add_heading(doc, "Takrorlash savollari", level=2)
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph(f"{i}. {q}")
        set_font(p.runs[0], size=12)
        p.paragraph_format.space_after = Pt(4)


def create_lesson_docx(topic: str, data: dict, output_dir: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data["title"])
    set_font(title_run, size=20, bold=True, color=(0, 70, 127))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(data["subtitle"])
    set_font(sub_run, size=12, color=(100, 100, 100))

    doc.add_paragraph()

    # Sections
    for section in data["sections"]:
        add_heading(doc, section["heading"], level=2)
        add_body(doc, section["body"])

    # Key terms
    doc.add_paragraph()
    add_key_terms(doc, data["key_terms"])

    # Questions
    doc.add_paragraph()
    add_questions(doc, data["questions"])

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run("5-sinf Tarix darsligi. Tugash muddati: 2026-yil iyun oxiri.")
    set_font(footer_run, size=10, color=(150, 150, 150))
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    safe_name = topic.replace(":", "").replace("/", "-").replace("'", "").strip()
    filename = output_dir / f"{safe_name}.docx"
    doc.save(filename)
    print(f"  Saved: {filename.name}")
    return filename


def main():
    output_dir = Path(__file__).parent / "lessons"
    output_dir.mkdir(exist_ok=True)

    print(f"Generating {len(LESSONS_CONTENT)} lesson files in: {output_dir}\n")

    saved = []
    for topic, data in LESSONS_CONTENT.items():
        try:
            path = create_lesson_docx(topic, data, output_dir)
            saved.append(path)
        except Exception as e:
            print(f"  ERROR {topic}: {e}")

    print(f"\nDone: {len(saved)} files saved to {output_dir}")


if __name__ == "__main__":
    main()
